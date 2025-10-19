from docx import Document
import os
import sys
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


def print_progress(message):
    print(f"\n> {message}")
    sys.stdout.flush()


class DocWriter:
    def __init__(self) -> None:
        self.doc = Document()
        # Configurar estilo del documento
        self.doc.styles['Normal'].font.name = 'Times New Roman'
        self.doc.styles['Normal'].font.size = Pt(12)

    def extract_clean_title(self, chapter_title):
        """Extrae solo el título limpio del capítulo sin marcos ni descripciones extensas"""
        # Si es prólogo o epílogo, dejar el título como está
        if chapter_title.lower().startswith(("prólogo", "epílogo")):
            return chapter_title
            
        # Si hay dos puntos, tomar solo la primera parte como título
        if ":" in chapter_title:
            title_part = chapter_title.split(":", 1)[0].strip()
            return title_part
            
        return chapter_title
    
    def clean_content(self, text):
        """
        Limpia el texto de posibles etiquetas, marcadores o contenido que no sea narrativo.
        
        NOTA: Esta función ahora usa el sistema unificado de limpieza de texto.
        Mantenida por compatibilidad con código existente.
        """
        from text_cleaning import clean_content as _clean_content
        return _clean_content(text, aggressive=True)

    def write_doc(self, book, chapter_dict, title, output_format='pdf', output_path='./docs'):
        try:
            print_progress("Iniciando creación del documento...")
            
            # Normalizar la ruta de salida
            if not os.path.isabs(output_path):
                output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), output_path.lstrip('./\\'))
            
            # Crear directorio de salida si no existe
            os.makedirs(output_path, exist_ok=True)
            
            # Añadir metadatos para mejor indexación
            core_properties = self.doc.core_properties
            core_properties.author = "AI Book Generator"
            core_properties.title = title
            core_properties.subject = "Novela generada con IA"
            core_properties.category = "Ficción"
            
            # Configurar márgenes para mejor legibilidad
            section = self.doc.sections[0]
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            
            # Título principal
            title_paragraph = self.doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar índice de contenido
            toc_heading = self.doc.add_heading("Índice", 1)
            toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Crear índice con todos los capítulos
            for chapter in book.keys():
                clean_title = self.extract_clean_title(chapter)
                toc_entry = self.doc.add_paragraph()
                toc_entry.add_run(clean_title)
                toc_entry.paragraph_format.left_indent = Inches(0.5)
                
            # Agregar salto de página después del índice
            self.doc.add_page_break()

            # Escribir cada capítulo
            for chapter, paragraphs_list in book.items():
                print_progress(f"Escribiendo capítulo: {chapter}")
                
                # Usar título limpio solo con el nombre del capítulo, sin descripciones extensas
                clean_title = self.extract_clean_title(chapter)
                
                # Título del capítulo
                chapter_heading = self.doc.add_heading(clean_title, 1)
                chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Combinar todo el contenido del capítulo para limpiarlo completamente
                full_chapter_content = "\n\n".join([p for p in paragraphs_list if p and p.strip()])
                
                # Limpiar el texto de todos los elementos no narrativos de una vez
                clean_full_content = self.clean_content(full_chapter_content)
                
                # Dividir en párrafos lógicos
                inner_paragraphs = clean_full_content.strip().split('\n\n')
                
                # Añadir cada párrafo al documento
                for inner_para in inner_paragraphs:
                    if inner_para.strip():
                        paragraph = self.doc.add_paragraph(inner_para.strip())
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragraph.paragraph_format.first_line_indent = Inches(0.25)
                        paragraph.paragraph_format.space_after = Pt(12)

                # Agregar salto de página entre capítulos
                if chapter != list(book.keys())[-1]:  # No añadir salto después del último
                    self.doc.add_page_break()

            # Generar un nombre de archivo único basado en el título
            safe_title = "".join([c if c.isalnum() else "_" for c in title]).lower()
            file_id = 1
            while os.path.exists(os.path.join(output_path, f"book_{safe_title[:30]}_{file_id}.docx")):
                file_id += 1
            
            base_filename = f"book_{safe_title[:30]}_{file_id}"
            
            # Guardar el documento DOCX
            docx_path = os.path.join(output_path, f"{base_filename}.docx")
            self.doc.save(docx_path)
            print_progress(f"Documento DOCX guardado como: {docx_path}")
            
            final_path = docx_path
            
            # Convertir a PDF (intentamos siempre por defecto)
            pdf_path = os.path.join(output_path, f"{base_filename}.pdf")
            print_progress("Intentando convertir a PDF...")
            
            # 1. Intentar usar LibreOffice para la conversión
            soffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                r"C:\Program Files\LibreOffice\app\program\soffice.exe",
            ]
            
            converted = False
            
            # Método 1: LibreOffice
            for soffice_path in soffice_paths:
                if os.path.exists(soffice_path):
                    try:
                        subprocess.run([
                            soffice_path,
                            '--headless',
                            '--convert-to',
                            'pdf',
                            '--outdir',
                            output_path,
                            docx_path
                        ], check=True, capture_output=True, timeout=30)
                        converted = True
                        print_progress(f"PDF generado exitosamente con LibreOffice: {pdf_path}")
                        final_path = pdf_path
                        break
                    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
                        print_progress(f"Error con LibreOffice: {str(e)}")
                        continue
            
            # Método 2: Intentar con unoconv si LibreOffice falló
            if not converted:
                try:
                    subprocess.run(['unoconv', '-f', 'pdf', '-o', output_path, docx_path], 
                                  check=True, capture_output=True, timeout=30)
                    converted = True
                    print_progress(f"PDF generado exitosamente con unoconv: {pdf_path}")
                    final_path = pdf_path
                except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError):
                    print_progress("Unoconv no disponible o falló")
            
            # Método 3: Intentar con docx2pdf si los anteriores fallaron
            if not converted:
                try:
                    # Intentar importar docx2pdf
                    import docx2pdf
                    docx2pdf.convert(docx_path, pdf_path)
                    converted = True
                    print_progress(f"PDF generado exitosamente con docx2pdf: {pdf_path}")
                    final_path = pdf_path
                except (ImportError, Exception) as e:
                    print_progress(f"docx2pdf no disponible o falló: {str(e)}")
            
            # Si ningún método funcionó pero el usuario quiere PDF
            if not converted and output_format.lower() == 'pdf':
                print_progress("No se pudo convertir a PDF. El archivo está disponible en formato DOCX.")
                print_progress("Para convertir a PDF manualmente, puedes abrir el archivo DOCX con Word/LibreOffice y guardar como PDF.")
            
            return final_path

        except Exception as e:
            print_progress(f"Error guardando el documento: {str(e)}")
            raise
